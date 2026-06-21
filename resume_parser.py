"""Extract plain text from resume files (PDF, DOCX)."""

from __future__ import annotations

from pathlib import Path

SUPPORTED_RESUME_EXT = {".pdf", ".docx"}
LEGACY_DOC_EXT = ".doc"


def extract_resume_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext == LEGACY_DOC_EXT:
        raise ValueError("Old .doc files are not supported. Save as PDF or DOCX and upload again.")
    raise ValueError("Unsupported resume format. Upload PDF or DOCX.")


def _extract_pdf(path: Path) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber is required for PDF resumes.") from exc

    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            cleaned = text.strip()
            if cleaned:
                parts.append(cleaned)
    text = "\n\n".join(parts).strip()
    if not text:
        raise ValueError("Could not read text from this PDF. Try a text-based PDF, not a scan.")
    return text


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for Word resumes.") from exc

    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Could not read text from this Word file.")
    return text
